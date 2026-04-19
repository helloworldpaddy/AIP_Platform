"""
Text extractors for local files: PDF, DOCX, TXT, MD.

The interface is intentionally simple:
    extract_text(path: Path) -> str
Callers handle chunking separately.

A small `StorageAdapter` protocol is included so callers can plug in
S3/GCS/etc without changing the ingestion pipeline. Local filesystem is
the only implementation provided.
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import Iterable, Protocol

from pypdf import PdfReader

log = logging.getLogger(__name__)


class StorageAdapter(Protocol):
    def list_files(self, root: str) -> Iterable[Path]: ...
    def open_bytes(self, path: Path) -> bytes: ...


class LocalFileStorage:
    def list_files(self, root: str) -> Iterable[Path]:
        base = Path(root)
        if base.is_file():
            yield base
            return
        for p in sorted(base.rglob("*")):
            if p.is_file():
                yield p

    def open_bytes(self, path: Path) -> bytes:
        return path.read_bytes()


# ---------------------------------------------------------------- extractors
def extract_text(path: Path) -> str:
    """Dispatch to an extractor based on file extension."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {ext} ({path})")


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as exc:  # pypdf occasionally raises on malformed PDFs
            log.warning("pdf.extract.page_failed", extra={"path": str(path), "err": str(exc)})
    return "\n\n".join(parts)


def _extract_docx(path: Path) -> str:
    # Imported lazily because python-docx pulls a fair bit at import time.
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(paragraphs)
